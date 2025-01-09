from flask import Flask, request, jsonify, send_file, Response, stream_with_context
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import time
from docx import Document
from pdf2docx import Converter
import win32com.client
import pythoncom
import uuid
import PyPDF2  # 添加PyPDF2用于获取页数
from pdf2docx.page.Pages import Pages
import json
import queue
import threading
from threading import Lock
import logging

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

# 配置上传文件夹
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 限制文件大小为16MB

ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}

# 创建一个全局的进度队列
progress_queues = {}

# 添加锁来保护队列操作
queue_lock = Lock()

def create_progress_queue():
    """创建一个新的进度队列并返回队列ID"""
    with queue_lock:
        queue_id = str(uuid.uuid4())
        progress_queues[queue_id] = queue.Queue()
        logger.info(f"创建队列 {queue_id}, 当前队列数: {len(progress_queues)}")
        logger.info(f"当前活动队列: {list(progress_queues.keys())}")
        return queue_id

def send_progress(queue_id, status):
    """发送进度信息到指定队列"""
    with queue_lock:
        if queue_id in progress_queues:
            logger.info(f"发送进度到队列 {queue_id}: {status}")
            progress_queues[queue_id].put(status)
        else:
            logger.error(f"尝试发送进度到不存在的队列 {queue_id}")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_unique_filename(original_filename):
    """生成唯一的文件名"""
    unique_id = str(uuid.uuid4())
    extension = os.path.splitext(original_filename)[1]
    return f"{unique_id}{extension}"

def convert_word_to_pdf(word_path, pdf_path):
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(word_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 represents PDF format
        doc.Close()
    finally:
        if word:
            word.Quit()
        pythoncom.CoUninitialize()

def safe_remove_file(file_path):
    """安全地删除文件，如果文件被占用则等待"""
    max_attempts = 5
    for attempt in range(max_attempts):
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
            break
        except PermissionError:
            if attempt < max_attempts - 1:
                time.sleep(1)  # 等待1秒后重试
            else:
                print(f"无法删除文件: {file_path}")

def get_pdf_page_count(pdf_path):
    """获取PDF文件的总页数"""
    with open(pdf_path, 'rb') as file:
        pdf = PyPDF2.PdfReader(file)
        return len(pdf.pages)

def optimize_pdf_to_word(pdf_path, word_path, queue_id):
    """优化的PDF转Word转换函数"""
    try:
        # 首先获取总页数
        total_pages = get_pdf_page_count(pdf_path)
        logger.info(f"PDF文件共 {total_pages} 页")
        send_progress(queue_id, {'status': f'PDF文件共 {total_pages} 页'})
        
        cv = Converter(pdf_path)
        
        def progress_callback(page):
            """进度回调函数"""
            current = page + 1
            status = f'正在转换第 {current}/{total_pages} 页'
            logger.info(status)
            send_progress(queue_id, {'status': status})
            return True
        
        # 设置转换参数
        cv.convert(word_path, start=0, end=None,
                  kwargs={
                      'min_section_height': 20,     # 最小段落高度
                      'connected_border_tolerance': 0.5,  # 连接边界容差
                      'line_overlap_threshold': 0.9,    # 行重叠阈值
                      'line_break_width_ratio': 0.1,   # 换行宽度比
                      'line_break_free_space_ratio': 0.1,  # 换行空白比
                      'line_separate_threshold': 5,    # 行分隔阈值
                      'line_space_threshold': 0.1,     # 行间距阈值
                      'paragraph_space_threshold': 0.3, # 段落间距阈值
                      'line_break_width_threshold': 0.5,# 换行宽度阈值
                  },
                  pages=range(total_pages),  # 明确指定页面范围
                  page_callback=progress_callback)
        
        logger.info("PDF转换完成")
        return True
    except Exception as e:
        logger.error(f"PDF转换失败: {str(e)}", exc_info=True)
        raise
    finally:
        cv.close()

@app.route('/progress/<queue_id>')
def get_progress(queue_id):
    """获取转换进度的SSE端点"""
    logger.info(f"收到进度请求，队列ID: {queue_id}")
    
    with queue_lock:
        logger.info(f"当前活动队列: {list(progress_queues.keys())}")
        if queue_id not in progress_queues:
            logger.error(f"无效的队列ID: {queue_id}")
            return Response(
                f"data: {json.dumps({'error': '无效的队列ID'})}\n\n",
                mimetype='text/event-stream',
                headers={
                    'Cache-Control': 'no-cache',
                    'Connection': 'keep-alive',
                    'X-Accel-Buffering': 'no'
                }
            )
    
    def generate():
        try:
            while True:
                with queue_lock:
                    if queue_id not in progress_queues:
                        logger.error(f"队列已被删除: {queue_id}")
                        yield f"data: {json.dumps({'error': '队列已被删除'})}\n\n"
                        break
                    q = progress_queues[queue_id]
                
                try:
                    logger.info(f"等待队列 {queue_id} 的进度更新...")
                    progress = q.get(timeout=30)
                    logger.info(f"队列 {queue_id} 的进度: {progress}")
                    
                    if isinstance(progress, dict):
                        if 'error' in progress:
                            logger.error(f"队列 {queue_id} 发生错误: {progress['error']}")
                            yield f"data: {json.dumps(progress)}\n\n"
                            break
                        if progress.get('status') == 'DONE':
                            logger.info(f"队列 {queue_id} 转换完成")
                            yield f"data: {json.dumps(progress)}\n\n"
                            break
                        if 'file_path' not in progress:
                            logger.info(f"发送进度更新: {progress}")
                            yield f"data: {json.dumps(progress)}\n\n"
                except queue.Empty:
                    logger.debug(f"队列 {queue_id} 等待超时，发送心跳包")
                    yield f"data: {json.dumps({'status': 'heartbeat'})}\n\n"
        except GeneratorExit:
            logger.info(f"客户端断开连接: {queue_id}")
        except Exception as e:
            logger.error(f"进度监听出错: {queue_id}, 错误: {str(e)}", exc_info=True)
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive',
            'X-Accel-Buffering': 'no'
        }
    )

@app.route('/convert/pdf-to-word', methods=['POST'])
def pdf_to_word():
    if 'file' not in request.files:
        return jsonify({'error': '没有文件上传'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '没有选择文件'}), 400
    
    if file and allowed_file(file.filename):
        try:
            # 生成文件路径
            filename = get_unique_filename(file.filename)
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            word_path = os.path.join(app.config['UPLOAD_FOLDER'], 
                                   os.path.splitext(filename)[0] + '.docx')
            
            # 保存上传的PDF文件
            file.save(pdf_path)
            
            # 转换PDF到Word
            cv = Converter(pdf_path)
            try:
                cv.convert(word_path)
            finally:
                cv.close()
            
            # 优化Word文档格式
            doc = Document(word_path)
            for paragraph in doc.paragraphs:
                paragraph.paragraph_format.space_before = 0
                paragraph.paragraph_format.space_after = 6
                paragraph.paragraph_format.line_spacing = 1.15
            doc.save(word_path)
            
            # 返回转换后的文件
            response = send_file(
                word_path,
                as_attachment=True,
                download_name=os.path.splitext(file.filename)[0] + '.docx'
            )
            
            # 注册清理函数
            response.call_on_close(lambda: safe_remove_file(pdf_path))
            response.call_on_close(lambda: safe_remove_file(word_path))
            
            return response
            
        except Exception as e:
            if 'pdf_path' in locals() and os.path.exists(pdf_path):
                safe_remove_file(pdf_path)
            if 'word_path' in locals() and os.path.exists(word_path):
                safe_remove_file(word_path)
            return jsonify({'error': f'转换失败: {str(e)}'}), 500
    
    return jsonify({'error': '不支持的文件类型'}), 400

@app.route('/convert/word-to-pdf', methods=['POST'])
def word_to_pdf():
    if 'file' not in request.files:
        return jsonify({'error': '没有文件上传'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '没有选择文件'}), 400
    
    if file and allowed_file(file.filename):
        # 使用唯一文件名
        filename = get_unique_filename(file.filename)
        word_path = os.path.join(os.path.abspath(app.config['UPLOAD_FOLDER']), filename)
        pdf_path = os.path.join(os.path.abspath(app.config['UPLOAD_FOLDER']), 
                               os.path.splitext(filename)[0] + '.pdf')
        
        try:
            file.save(word_path)
            
            # 转换Word到PDF
            convert_word_to_pdf(word_path, pdf_path)
            
            # 返回转换后的文件
            response = send_file(pdf_path, as_attachment=True,
                               download_name=os.path.splitext(file.filename)[0] + '.pdf')
            
            # 使用response.call_on_close注册清理函数
            response.call_on_close(lambda: safe_remove_file(word_path))
            response.call_on_close(lambda: safe_remove_file(pdf_path))
            
            return response
            
        except Exception as e:
            safe_remove_file(word_path)
            safe_remove_file(pdf_path)
            return jsonify({'error': f'转换失败: {str(e)}'}), 500
    
    return jsonify({'error': '不支持的文件类型'}), 400

# 添加文件下载路由
@app.route('/download/<queue_id>')
def download_file(queue_id):
    """下载转换后的文件"""
    if queue_id not in progress_queues:
        return jsonify({'error': '文件不存在或已过期'}), 404
    
    try:
        q = progress_queues[queue_id]
        # 检查队列中是否有数据
        if q.empty():
            return jsonify({'error': '文件还未准备好'}), 400
        
        # 获取但不删除队列中的数据
        file_info = q.queue[0]  # 直接访问队列中的第一个元素
        
        if isinstance(file_info, dict) and 'file_path' in file_info:
            file_path = file_info['file_path']
            if not os.path.exists(file_path):
                return jsonify({'error': '文件已被删除'}), 404
                
            return send_file(
                file_path,
                as_attachment=True,
                download_name=os.path.splitext(file_info['original_name'])[0] + '.docx'
            )
        return jsonify({'error': '文件格式错误'}), 400
        
    except Exception as e:
        return jsonify({'error': f'下载失败: {str(e)}'}), 500

def process_conversion(temp_path, original_filename, queue_id):
    """处理文件转换的函数"""
    pdf_path = None
    word_path = None
    
    logger.info(f"开始处理转换，队列ID: {queue_id}")
    logger.info(f"原始文件名: {original_filename}")
    
    try:
        # 生成最终的文件路径
        filename = get_unique_filename(original_filename)
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        word_path = os.path.join(app.config['UPLOAD_FOLDER'], 
                                os.path.splitext(filename)[0] + '.docx')
        
        logger.info(f"移动临时文件 {temp_path} 到 {pdf_path}")
        os.rename(temp_path, pdf_path)
        send_progress(queue_id, {'status': '文件准备完成，开始转换...'})
        
        # 使用优化的转换函数
        logger.info("开始PDF转Word转换...")
        optimize_pdf_to_word(pdf_path, word_path, queue_id)
        logger.info("PDF转Word转换完成")
        
        # 进一步优化Word文档格式
        logger.info("开始优化Word文档格式...")
        send_progress(queue_id, {'status': '正在优化文档格式...'})
        doc = Document(word_path)
        
        # 调整段落格式
        paragraph_count = len(doc.paragraphs)
        logger.info(f"正在处理 {paragraph_count} 个段落...")
        for i, paragraph in enumerate(doc.paragraphs, 1):
            if i % 100 == 0:  # 每处理100个段落记录一次日志
                logger.info(f"已处理 {i}/{paragraph_count} 个段落")
                send_progress(queue_id, {'status': f'正在优化格式 ({i}/{paragraph_count})'})
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 6
            paragraph.paragraph_format.line_spacing = 1.15
        
        logger.info("正在保存优化后的Word文档...")
        send_progress(queue_id, {'status': '正在保存文档...'})
        doc.save(word_path)
        logger.info("Word文档保存完成")
        
        # 发送文件路径和完成状态
        with queue_lock:
            if queue_id in progress_queues:
                logger.info(f"将转换后的文件信息添加到队列: {queue_id}")
                progress_queues[queue_id].put({
                    'file_path': word_path,
                    'original_name': original_filename
                })
                send_progress(queue_id, {'status': 'DONE'})
        
        logger.info("等待5秒确保文件被下载...")
        time.sleep(5)
        logger.info("转换流程完成")
        
    except Exception as e:
        logger.error(f"转换过程出错: {str(e)}", exc_info=True)
        send_progress(queue_id, {'error': str(e)})
    finally:
        logger.info("开始清理临时文件...")
        # 清理文件
        if temp_path and os.path.exists(temp_path):
            logger.info(f"清理临时文件: {temp_path}")
            safe_remove_file(temp_path)
        if pdf_path and os.path.exists(pdf_path):
            logger.info(f"清理PDF文件: {pdf_path}")
            safe_remove_file(pdf_path)
        if word_path and os.path.exists(word_path):
            logger.info(f"清理Word文件: {word_path}")
            safe_remove_file(word_path)
        # 最后再删除队列
        with queue_lock:
            if queue_id in progress_queues:
                logger.info(f"删除队列: {queue_id}")
                del progress_queues[queue_id]
        logger.info("清理完成")

if __name__ == '__main__':
    try:
        logger.info("正在启动服务器...")
        logger.info(f"上传文件夹路径: {os.path.abspath(UPLOAD_FOLDER)}")
        logger.info("服务器配置:")
        logger.info(f"- 最大文件大小: {app.config['MAX_CONTENT_LENGTH'] / 1024 / 1024}MB")
        logger.info(f"- 允许的文件类型: {ALLOWED_EXTENSIONS}")
        logger.info("- CORS已启用")
        
        # 确保上传文件夹存在
        if not os.path.exists(UPLOAD_FOLDER):
            os.makedirs(UPLOAD_FOLDER)
            logger.info(f"创建上传文件夹: {UPLOAD_FOLDER}")
        
        # 清理可能存在的临时文件
        for file in os.listdir(UPLOAD_FOLDER):
            try:
                file_path = os.path.join(UPLOAD_FOLDER, file)
                if os.path.isfile(file_path):
                    os.remove(file_path)
                    logger.info(f"清理临时文件: {file}")
            except Exception as e:
                logger.warning(f"清理文件失败: {file}, 错误: {str(e)}")
        
        # 启动服务器
        logger.info("服务器启动完成，监听端口: 5000")
        app.run(debug=True, port=5000, use_reloader=False)
    except Exception as e:
        logger.error(f"服务器启动失败: {str(e)}")
        raise 